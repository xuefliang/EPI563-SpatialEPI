from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy

def split_word_by_pages(input_file, output_dir, pages_list):
    """
    按照指定页码拆分Word文档
    
    Args:
        input_file: 输入Word文件路径
        output_dir: 输出目录
        pages_list: 页码列表，如 [[1,5], [6,10], [11,15]]
                   表示分别拆分为 1-5页、6-10页、11-15页
    """
    import os
    from docx import Document
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 打开原始文档
    doc = Document(input_file)
    
    # 遍历每个页码范围
    for idx, (start_page, end_page) in enumerate(pages_list):
        # 创建新文档
        new_doc = Document()
        
        # 复制样式
        for style in doc.styles:
            try:
                new_doc.styles.add_style(style.name, style.type)
            except:
                pass
        
        # 根据段落数估算页码（这是个近似方法）
        # 实际上python-docx无法准确获取页码，需要其他方法
        para_count = len(doc.paragraphs)
        
        # 这里需要更精确的页码识别方法
        start_para = int(para_count * (start_page - 1) / estimate_total_pages(doc))
        end_para = int(para_count * end_page / estimate_total_pages(doc))
        
        # 复制段落
        for para in doc.paragraphs[start_para:end_para]:
            new_para = new_doc.add_paragraph(para.text)
            new_para.style = para.style
        
        # 保存新文档
        output_file = os.path.join(output_dir, f'split_{idx+1}.docx')
        new_doc.save(output_file)
        print(f"已生成: {output_file}")

def estimate_total_pages(doc):
    """估算文档总页数"""
    return max(1, len(doc.paragraphs) // 30)  # 粗略估算